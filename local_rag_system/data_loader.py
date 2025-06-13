"""数据加载和处理模块"""

import os
import hashlib
from typing import Tuple, List, Optional
from config import (DATA_FOLDER, SUPPORTED_EXTENSIONS, PRESERVE_URLS, PRESERVE_EMAIL, 
                   PRESERVE_SPECIAL_CHARS, MINIMAL_FILTERING, PRESERVE_ORIGINAL_TEXT,
                   MAX_CHUNK_SIZE, MIN_CHUNK_SIZE, OVERLAP_SIZE, MERGE_SHORT_LINES, 
                   PRESERVE_PARAGRAPHS, SENTENCE_SPLIT)
from docx import Document


class DataLoader:
    """数据加载器"""
    
    def __init__(self, folder_path: str = DATA_FOLDER):
        self.folder_path = folder_path
    
    def load_from_folder(self) -> Tuple[Optional[str], List[str]]:
        """从文件夹加载所有支持的文件"""
        if not os.path.exists(self.folder_path):
            print(f"数据文件夹 {self.folder_path} 不存在")
            return None, []
        
        all_content = ""
        file_list = []
        documents = []
        
        try:
            for filename in os.listdir(self.folder_path):
                file_path = os.path.join(self.folder_path, filename)
                
                if self._is_supported_file(file_path, filename):
                    content, docs = self._process_file(file_path, filename)
                    if content and docs:
                        all_content += content + "\n"
                        documents.extend(docs)
                        file_list.append(filename)
                        print(f"加载文件: {filename} ({len(docs)} 个文档)")
            
            if not documents:
                print("未找到有效的数据文件")
                return None, []
                
            print(f"成功加载 {len(file_list)} 个文件，共 {len(documents)} 个文档")
            return all_content, documents
            
        except Exception as e:
            print(f"读取数据文件夹失败: {e}")
            return None, []
    
    def _is_supported_file(self, file_path: str, filename: str) -> bool:
        """检查文件是否为支持的格式"""
        if not os.path.isfile(file_path):
            return False
        
        # 跳过临时文件和系统文件
        if self._is_temp_or_system_file(filename):
            print(f"跳过临时文件: {filename}")
            return False
        
        file_ext = os.path.splitext(filename)[1].lower()
        if file_ext not in SUPPORTED_EXTENSIONS:
            print(f"跳过不支持的文件格式: {filename}")
            return False
        
        return True
    
    def _is_temp_or_system_file(self, filename: str) -> bool:
        """检查是否为临时文件或系统文件"""
        # Word临时文件通常以~$开头
        if filename.startswith('~$'):
            return True
        
        # Excel临时文件
        if filename.startswith('~') and filename.endswith('.tmp'):
            return True
        
        # 隐藏文件（以.开头）
        if filename.startswith('.'):
            return True
        
        # 其他常见的临时文件模式
        temp_patterns = [
            '.tmp', '.temp', '.bak', '.swp', '.lock',
            'Thumbs.db', 'Desktop.ini', '.DS_Store'
        ]
        
        for pattern in temp_patterns:
            if filename.lower().endswith(pattern.lower()) or pattern.lower() in filename.lower():
                return True
        
        return False
    
    def _process_file(self, file_path: str, filename: str) -> Tuple[Optional[str], List[str]]:
        """处理单个文件"""
        try:
            file_ext = os.path.splitext(filename)[1].lower()
            
            if file_ext == '.docx':
                content = self._read_docx_file(file_path)
            else:
                content = self._read_text_file(file_path)
            
            if content is None:
                print(f"文件 {filename} 内容为空或读取失败")
                return None, []
            
            # 根据文件类型处理内容
            if file_ext == '.docx':
                documents = self._process_docx_content(content, filename)
            else:
                documents = self._process_text_content(content, filename, file_ext)
            
            print(f"文件 {filename} 处理完成，提取了 {len(documents)} 个文档片段")
            return content, documents
            
        except Exception as e:
            print(f"读取文件 {filename} 失败: {e}")
            return None, []
    
    def _process_text_content(self, content: str, filename: str, file_ext: str) -> List[str]:
        """处理文本内容，使用细粒度切割"""
        lines = [line.strip() for line in content.split('\n') if line.strip()]
        documents = []
        
        if not lines:
            return documents
        
        # 跳过表头
        start_idx = 1 if lines and self._is_header_line(lines[0]) else 0
        content_lines = lines[start_idx:]
        
        # 细粒度处理：每行或每个小块单独处理
        for line in content_lines:
            if len(line) <= MAX_CHUNK_SIZE:
                # 短行直接作为一个文档
                if len(line) >= MIN_CHUNK_SIZE:
                    doc_content = f"[文件: {filename}] {line}"
                    documents.append(doc_content)
            else:
                # 长行需要进一步切割
                chunks = self._split_long_text(line, filename)
                documents.extend(chunks)
        
        return documents
    
    def _process_docx_content(self, content: str, filename: str) -> List[str]:
        """处理docx内容，使用细粒度切割"""
        documents = []
        paragraphs = [p.strip() for p in content.split('\n') if p.strip()]
        
        if not paragraphs:
            return documents
        
        for paragraph in paragraphs:
            cleaned_paragraph = self._clean_text(paragraph)
            if not cleaned_paragraph or len(cleaned_paragraph) < MIN_CHUNK_SIZE:
                continue
            
            if SENTENCE_SPLIT:
                # 句子级别切割
                sentences = self._split_into_sentences(cleaned_paragraph)
                for sentence in sentences:
                    if len(sentence.strip()) >= MIN_CHUNK_SIZE:
                        # 进一步检查是否需要切割
                        if len(sentence) <= MAX_CHUNK_SIZE:
                            doc_content = f"[文件: {filename}] {sentence.strip()}"
                            documents.append(doc_content)
                        else:
                            # 超长句子继续切割
                            chunks = self._split_long_text(sentence.strip(), filename)
                            documents.extend(chunks)
            else:
                # 段落级别切割
                if len(cleaned_paragraph) <= MAX_CHUNK_SIZE:
                    doc_content = f"[文件: {filename}] {cleaned_paragraph}"
                    documents.append(doc_content)
                else:
                    # 切割长段落
                    chunks = self._split_long_text(cleaned_paragraph, filename)
                    documents.extend(chunks)
        
        return documents
    
    def _split_into_sentences(self, text: str) -> List[str]:
        """将文本切割成句子"""
        import re
        
        # 中文句子分割
        sentences = re.split(r'[。！？；]', text)
        
        # 英文句子分割
        result = []
        for sentence in sentences:
            if '.' in sentence and len(sentence) > 50:
                # 进一步按英文句号分割
                sub_sentences = re.split(r'\.(?=\s+[A-Z])', sentence)
                for sub in sub_sentences:
                    if sub.strip():
                        result.append(sub.strip() + ('。' if not sub.endswith('.') else ''))
            else:
                if sentence.strip():
                    result.append(sentence.strip() + '。')
        
        return [s for s in result if len(s.strip()) > 5]
    
    def _split_long_text(self, text: str, filename: str) -> List[str]:
        """切割超长文本为小块"""
        chunks = []
        
        # 按字符数切割
        start = 0
        while start < len(text):
            end = start + MAX_CHUNK_SIZE
            
            # 尝试在合适位置断开（标点符号）
            if end < len(text):
                # 寻找最近的标点符号
                punctuation_pos = -1
                for i in range(end, max(start + MIN_CHUNK_SIZE, end - 100), -1):
                    if text[i] in '。！？；，、':
                        punctuation_pos = i + 1
                        break
                
                if punctuation_pos > 0:
                    end = punctuation_pos
            
            chunk = text[start:end].strip()
            if len(chunk) >= MIN_CHUNK_SIZE:
                doc_content = f"[文件: {filename}] {chunk}"
                chunks.append(doc_content)
            
            # 添加重叠
            start = end - OVERLAP_SIZE if OVERLAP_SIZE > 0 and end < len(text) else end
        
        return chunks

    def _clean_text(self, text: str) -> str:
        """最小化清理文本内容，保持原始信息完整性"""
        import re
        
        if PRESERVE_ORIGINAL_TEXT:
            # 只进行最基本的清理
            # 移除不可见的控制字符，但保留所有可见字符
            text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
            
            # 只替换多个连续空白为单个空格
            text = re.sub(r'\s+', ' ', text).strip()
            
            return text
        
        # 如果不使用原始文本保护，则使用之前的逻辑
        return self._clean_text_with_links(text)
    
    def _clean_text_with_links(self, text: str) -> str:
        """保守的文本清理，保护重要信息"""
        import re
        
        if MINIMAL_FILTERING:
            # 最小化过滤模式
            # 只移除明显的控制字符和不可打印字符
            text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
            
            # 保留几乎所有可见字符，只规范化空白
            text = re.sub(r'\s+', ' ', text).strip()
            
            return text
        
        # 原有的链接保护逻辑（保留作为备选）
        url_placeholders = {}
        url_pattern = r'https?://[^\s<>"\']+|ftp://[^\s<>"\']+|www\.[^\s<>"\']+\.[a-z]{2,}[^\s<>"\']+'
        
        def replace_url(match):
            url = match.group(0)
            placeholder = f"URLPROTECT{len(url_placeholders)}URLPROTECT"
            url_placeholders[placeholder] = url
            return placeholder
        
        text = re.sub(url_pattern, replace_url, text, flags=re.IGNORECASE)
        
        # 移除控制字符
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f]', '', text)
        
        # 非常宽松的字符过滤，几乎保留所有字符
        text = re.sub(r'[\x00-\x1f\x7f-\x9f]', ' ', text)
        
        # 替换多个空白为单个空格
        text = re.sub(r'\s+', ' ', text).strip()
        
        # 恢复保护的内容
        for placeholder, original in url_placeholders.items():
            text = text.replace(placeholder, original)
        
        return text

    def _read_text_file(self, file_path: str) -> Optional[str]:
        """读取文本文件"""
        try:
            with open(file_path, 'r', encoding='utf-8') as f:
                return f.read()
        except UnicodeDecodeError:
            # 尝试其他编码
            try:
                with open(file_path, 'r', encoding='gbk') as f:
                    return f.read()
            except:
                return None
        except Exception:
            return None
    
    def _read_docx_file(self, file_path: str) -> Optional[str]:
        """读取docx文件"""
        try:
            # 检查文件是否存在且不是临时文件
            if not os.path.exists(file_path):
                print(f"文件不存在: {file_path}")
                return None
            
            # 检查文件大小
            file_size = os.path.getsize(file_path)
            if file_size == 0:
                print(f"文件为空: {file_path}")
                return None
            
            print(f"开始处理docx文件: {os.path.basename(file_path)} (大小: {file_size} 字节)")
            
            doc = Document(file_path)
            content_lines = []
            
            print(f"文档总段落数: {len(doc.paragraphs)}")
            print(f"文档总表格数: {len(doc.tables)}")
            
            # 读取所有段落（包括所有页面）
            paragraph_count = 0
            for i, paragraph in enumerate(doc.paragraphs):
                text = paragraph.text.strip()
                if text:
                    # 使用最小化清理
                    cleaned_text = self._clean_text(text)
                    if cleaned_text and len(cleaned_text) > 1:  # 降低最小长度要求
                        content_lines.append(cleaned_text)
                        paragraph_count += 1
                        # 调试：显示前3个段落内容
                        if i < 3:
                            print(f"段落 {i+1}: {cleaned_text[:80]}...")
            
            print(f"有效段落数: {paragraph_count}")
            
            # 读取所有表格
            table_count = 0
            for table_idx, table in enumerate(doc.tables):
                print(f"处理表格 {table_idx + 1}，行数: {len(table.rows)}")
                for row_idx, row in enumerate(table.rows):
                    row_data = []
                    for cell in row.cells:
                        cell_text = self._clean_text(cell.text.strip())
                        if cell_text:  # 移除长度限制
                            row_data.append(cell_text)
                    if row_data:
                        if row_idx == 0:
                            content_lines.append(f"表格{table_idx + 1}数据：{' | '.join(row_data)}")
                        else:
                            content_lines.append(f"{' | '.join(row_data)}")
                        table_count += 1
            
            if table_count > 0:
                print(f"提取了 {table_count} 个表格行")
            
            final_content = '\n'.join(content_lines)
            print(f"docx文件总内容长度: {len(final_content)} 字符")
            
            return final_content
            
        except Exception as e:
            print(f"读取docx文件失败: {e}")
            # 检查是否是文件访问权限问题
            if "Permission denied" in str(e) or "PackageNotFoundError" in str(e):
                print(f"可能是文件被占用或损坏，跳过文件: {os.path.basename(file_path)}")
            else:
                import traceback
                traceback.print_exc()
            return None
    
    def _is_header_line(self, line: str) -> bool:
        """检查是否为表头行"""
        return any(sep in line for sep in ['\t', ',', '|'])
    
    def get_folder_hash(self) -> str:
        """计算文件夹内容的哈希值"""
        if not os.path.exists(self.folder_path):
            return ""
        
        hash_md5 = hashlib.md5()
        
        # 获取所有文件的修改时间和内容
        for filename in sorted(os.listdir(self.folder_path)):
            file_path = os.path.join(self.folder_path, filename)
            if os.path.isfile(file_path):
                # 添加文件名和修改时间到哈希
                hash_md5.update(filename.encode('utf-8'))
                hash_md5.update(str(os.path.getmtime(file_path)).encode('utf-8'))
                
                # 添加文件内容到哈希
                try:
                    with open(file_path, 'rb') as f:
                        for chunk in iter(lambda: f.read(4096), b""):
                            hash_md5.update(chunk)
                except:
                    continue
        
        return hash_md5.hexdigest()

    def _is_likely_title(self, text: str) -> bool:
        """判断文本是否可能是标题"""
        # 简单的标题判断逻辑
        if len(text) < 50 and ('第' in text and '章' in text or 
                              '第' in text and '节' in text or
                              text.endswith('：') or
                              text.count('、') == 0):
            return True
        return False
